#!/usr/bin/env python3
"""
Generate a well-formatted Word document explaining the proposed 
PUF-based Decoupled Distributed Authentication Scheme.
Includes sequence diagrams for each phase.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── matplotlib for sequence diagrams ──
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ═══════════════════════════════════════════════════════════════════
# HELPER: set cell shading
# ═══════════════════════════════════════════════════════════════════
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Add borders to table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════════════
# SEQUENCE DIAGRAM GENERATORS
# ═══════════════════════════════════════════════════════════════════

def draw_enrollment_diagram(filepath):
    """Draw Node Enrollment Phase sequence diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 11))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis('off')
    
    # Title
    ax.text(5, 13.5, 'Phase 1: Node Enrollment', fontsize=16, fontweight='bold',
            ha='center', va='center', color='#1a3c6e')
    
    # Actors
    actors = {'D': 2.5, 'AS': 7.5}
    actor_colors = {'D': '#4CAF50', 'AS': '#2196F3'}
    
    for name, x in actors.items():
        label = 'Device Node (D)' if name == 'D' else 'Auth Server (AS)'
        rect = FancyBboxPatch((x-1.2, 12.5), 2.4, 0.7, boxstyle="round,pad=0.1",
                               facecolor=actor_colors[name], edgecolor='black', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, 12.85, label, ha='center', va='center', fontsize=10,
                fontweight='bold', color='white')
        # Lifeline
        ax.plot([x, x], [12.5, 1.0], '--', color='gray', linewidth=1, alpha=0.5)
    
    y = 12.0
    step = 0
    
    def msg(y, x1, x2, label, note=None, color='#333'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.5))
        mid = (x1 + x2) / 2
        ax.text(mid, y + 0.22, label, ha='center', va='bottom', fontsize=8.5,
                fontweight='bold', color=color)
        if note:
            ax.text(mid, y - 0.22, note, ha='center', va='top', fontsize=7,
                    color='#555', style='italic')
    
    def action(y, x, text, color='#FFF9C4'):
        rect = FancyBboxPatch((x-1.5, y-0.3), 3.0, 0.5, boxstyle="round,pad=0.05",
                               facecolor=color, edgecolor='#888', linewidth=0.8)
        ax.add_patch(rect)
        ax.text(x, y-0.05, text, ha='center', va='center', fontsize=7.5)
    
    # Step 1: D sends registration request
    y -= 0.5
    msg(y, 2.5, 7.5, 'Reg-0: SE(K_AS_D, [ID_D | pad])',
        'Registration request (16 B encrypted)', '#1a3c6e')
    
    # Step 2: AS generates c_d, m_d
    y -= 0.8
    action(y, 7.5, 'Generate c_D ∈ C_D, m_D (random)', '#E3F2FD')
    
    # Step 3: AS replies with c_d, m_d
    y -= 0.7
    msg(y, 7.5, 2.5, 'SE(K_AS_D, [c_D | m_D | pad])',
        'Challenge + session random (48 B encrypted)', '#1a3c6e')
    
    # Step 4: D computes PUF response
    y -= 0.8
    action(y, 2.5, 'R_D = PUF_D(c_D)', '#E8F5E9')
    
    # Step 5: D generates y_d and computes Y_dH
    y -= 0.7
    action(y, 2.5, 'Generate y_D; Y_D^H = H(y_D)', '#E8F5E9')
    
    # Step 6: D picks c_as_d
    y -= 0.7
    action(y, 2.5, 'Pick c_AS-D ∈ C_AS-D', '#E8F5E9')
    
    # Step 7: D sends Reg-1
    y -= 0.7
    msg(y, 2.5, 7.5, 'Reg-1: SE(K_AS_D, [ID_D | Y_D^H | R_D | c_AS-D])',
        '48 B encrypted payload', '#1a3c6e')
    
    # Step 8: AS accumulates
    y -= 0.8
    action(y, 7.5, 'T_acc = T_acc & Y_D^H', '#E3F2FD')
    
    # Step 9: AS computes phi
    y -= 0.7
    action(y, 7.5, 'R_AS = PUF_AS(c_AS-D)', '#E3F2FD')
    
    y -= 0.7
    action(y, 7.5, 'Φ_AS-D = R_AS ⊕ R_D', '#E3F2FD')
    
    # Step 10: AS computes initial PID
    y -= 0.7
    action(y, 7.5, 'PID_curr = H(ID_D || m_D)', '#E3F2FD')
    
    # Storage boxes
    y -= 0.9
    rect_d = FancyBboxPatch((0.5, y-0.35), 3.8, 0.65, boxstyle="round,pad=0.05",
                             facecolor='#E8F5E9', edgecolor='#4CAF50', linewidth=1.2)
    ax.add_patch(rect_d)
    ax.text(2.5, y+0.05, 'D stores: y_D, c_D, m_D', ha='center', va='center',
            fontsize=8, fontweight='bold', color='#2E7D32')
    
    rect_as = FancyBboxPatch((5.3, y-0.55), 4.2, 1.0, boxstyle="round,pad=0.05",
                              facecolor='#E3F2FD', edgecolor='#2196F3', linewidth=1.2)
    ax.add_patch(rect_as)
    ax.text(7.4, y+0.1, 'AS stores: Φ_AS-D, c_AS-D, m_D,', ha='center', va='center',
            fontsize=8, fontweight='bold', color='#1565C0')
    ax.text(7.4, y-0.2, 'PID_curr, T_acc (updated)', ha='center', va='center',
            fontsize=8, fontweight='bold', color='#1565C0')
    
    # Note
    ax.text(5, 1.2, '※ Enrollment is executed over a secure (encrypted) channel',
            ha='center', va='center', fontsize=8.5, style='italic', color='#666')
    
    plt.tight_layout()
    plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def draw_authentication_diagram(filepath):
    """Draw Node Authentication Phase sequence diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 10))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 13)
    ax.axis('off')
    
    ax.text(5, 12.5, 'Phase 2: Node Authentication', fontsize=16, fontweight='bold',
            ha='center', va='center', color='#1a3c6e')
    
    actors = {'D': 2.5, 'AS': 7.5}
    actor_colors = {'D': '#4CAF50', 'AS': '#2196F3'}
    
    for name, x in actors.items():
        label = 'Device Node (D)' if name == 'D' else 'Auth Server (AS)'
        rect = FancyBboxPatch((x-1.2, 11.7), 2.4, 0.7, boxstyle="round,pad=0.1",
                               facecolor=actor_colors[name], edgecolor='black', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, 12.05, label, ha='center', va='center', fontsize=10,
                fontweight='bold', color='white')
        ax.plot([x, x], [11.7, 1.5], '--', color='gray', linewidth=1, alpha=0.5)
    
    y = 11.2
    
    def msg(y, x1, x2, label, note=None, color='#333'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.5))
        mid = (x1 + x2) / 2
        ax.text(mid, y + 0.22, label, ha='center', va='bottom', fontsize=8.5,
                fontweight='bold', color=color)
        if note:
            ax.text(mid, y - 0.22, note, ha='center', va='top', fontsize=7,
                    color='#555', style='italic')
    
    def action(y, x, text, color='#FFF9C4'):
        rect = FancyBboxPatch((x-1.8, y-0.3), 3.6, 0.5, boxstyle="round,pad=0.05",
                               facecolor=color, edgecolor='#888', linewidth=0.8)
        ax.add_patch(rect)
        ax.text(x, y-0.05, text, ha='center', va='center', fontsize=7.5)
    
    # D computes mask and sends auth
    action(y, 2.5, 'R_D = PUF_D(c_D)', '#E8F5E9')
    y -= 0.65
    action(y, 2.5, 'PID = H(ID_D || m_D)', '#E8F5E9')
    y -= 0.65
    action(y, 2.5, 'Y_D^H = H(y_D)', '#E8F5E9')
    y -= 0.65
    action(y, 2.5, 'mask = H(R_D || m_D || PID || ts₁)', '#E8F5E9')
    y -= 0.65
    action(y, 2.5, 'Y_ASD = Y_D^H ⊕ mask', '#E8F5E9')
    
    y -= 0.7
    msg(y, 2.5, 7.5, 'PID | Y_ASD | ts₁',
        'Authentication request (65 B plaintext)', '#1a3c6e')
    
    # AS verifies
    y -= 0.7
    action(y, 7.5, 'Search PID in {PID_curr, PID_old}', '#E3F2FD')
    y -= 0.65
    action(y, 7.5, 'Freshness check: ts₁ > last_ts₁', '#E3F2FD')
    y -= 0.65
    action(y, 7.5, 'R_AS=PUF_AS(c_AS-D); R_D=Φ_AS-D⊕R_AS', '#E3F2FD')
    y -= 0.65
    action(y, 7.5, 'mask = H(R_D || m_D || PID || ts₁)', '#E3F2FD')
    y -= 0.65
    action(y, 7.5, 'Y_D^H = Y_ASD ⊕ mask', '#E3F2FD')
    y -= 0.65
    action(y, 7.5, 'Verify: T_acc & Y_D^H == T_acc', '#E3F2FD')
    
    # Result
    y -= 0.7
    rect = FancyBboxPatch((4.0, y-0.25), 3.0, 0.5, boxstyle="round,pad=0.1",
                           facecolor='#C8E6C9', edgecolor='#4CAF50', linewidth=1.5)
    ax.add_patch(rect)
    ax.text(5.5, y, '✓ D is Authenticated', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#2E7D32')
    
    ax.text(5, 1.8, '※ Authentication uses open channel — security provided by PUF + hash masking',
            ha='center', va='center', fontsize=8.5, style='italic', color='#666')
    
    plt.tight_layout()
    plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def draw_key_exchange_diagram(filepath):
    """Draw Key Exchange Phase sequence diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(12, 12))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 15)
    ax.axis('off')
    
    ax.text(6, 14.5, 'Phase 3: Session Key Exchange & Mask Update',
            fontsize=16, fontweight='bold', ha='center', va='center', color='#1a3c6e')
    
    actors = {'D': 2.0, 'AS': 6.0, 'GW': 10.0}
    actor_colors = {'D': '#4CAF50', 'AS': '#2196F3', 'GW': '#FF9800'}
    actor_labels = {'D': 'Device (D)', 'AS': 'Auth Server (AS)', 'GW': 'Gateway (GW)'}
    
    for name, x in actors.items():
        rect = FancyBboxPatch((x-1.2, 13.5), 2.4, 0.7, boxstyle="round,pad=0.1",
                               facecolor=actor_colors[name], edgecolor='black', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, 13.85, actor_labels[name], ha='center', va='center', fontsize=10,
                fontweight='bold', color='white')
        ax.plot([x, x], [13.5, 1.0], '--', color='gray', linewidth=1, alpha=0.5)
    
    y = 13.0
    
    def msg(y, x1, x2, label, note=None, color='#1a3c6e'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.5))
        mid = (x1 + x2) / 2
        ax.text(mid, y + 0.2, label, ha='center', va='bottom', fontsize=8,
                fontweight='bold', color=color)
        if note:
            ax.text(mid, y - 0.2, note, ha='center', va='top', fontsize=6.5,
                    color='#555', style='italic')
    
    def action(y, x, text, color='#FFF9C4', width=3.0):
        rect = FancyBboxPatch((x-width/2, y-0.25), width, 0.45, boxstyle="round,pad=0.05",
                               facecolor=color, edgecolor='#888', linewidth=0.8)
        ax.add_patch(rect)
        ax.text(x, y-0.02, text, ha='center', va='center', fontsize=7)
    
    # AS generates key material
    action(y, 6.0, 'Generate n₁ (32B random)', '#E3F2FD')
    y -= 0.55
    action(y, 6.0, 'm_new = H(n₁)', '#E3F2FD')
    y -= 0.55
    action(y, 6.0, 'K_GW-D = H(R_D || m_new)', '#E3F2FD')
    
    # AS computes masked m_new
    y -= 0.55
    action(y, 6.0, 'mH_mask = H(Y_D^H||m_D||R_D||ID_AS||PID||ts₂)', '#E3F2FD', width=3.8)
    y -= 0.55
    action(y, 6.0, 'm_H = m_new ⊕ mH_mask', '#E3F2FD')
    
    # AS sends masked m_new to D
    y -= 0.6
    msg(y, 6.0, 2.0, 'ACK | m_H | ts₂',
        'Key exchange reply (34 B)', '#1a3c6e')
    
    # AS builds and sends token to GW
    y -= 0.65
    action(y, 6.0, 'ts_auth = clock()', '#E3F2FD')
    y -= 0.55
    action(y, 6.0, 'token = SE(K_GW-AS, [ID_D|ID_AS|ts_auth|K_GW-D])', '#E3F2FD', width=3.8)
    
    y -= 0.6
    msg(y, 6.0, 10.0, 'new_PID | ID_AS | enc_token',
        'Auth token (81 B) via CoAP POST', '#E65100')
    
    # D recovers m_new
    y -= 0.7
    action(y, 2.0, 'Recover m_new = m_H ⊕ mH_mask', '#E8F5E9')
    y -= 0.55
    action(y, 2.0, 'K_GW-D = H(R_D || m_new)', '#E8F5E9')
    y -= 0.55
    action(y, 2.0, 'Update: m_D ← m_new', '#E8F5E9')
    y -= 0.55
    action(y, 2.0, 'PID ← H(ID_D || m_new)', '#E8F5E9')
    
    # GW processes token
    y_gw = 9.2
    action(y_gw, 10.0, 'Decrypt token with K_GW-AS', '#FFF3E0')
    y_gw -= 0.55
    action(y_gw, 10.0, 'Verify ts_auth freshness', '#FFF3E0')
    y_gw -= 0.55
    action(y_gw, 10.0, 'Store session: {PID, K_GW-D}', '#FFF3E0')
    
    # Data communication
    y = min(y, y_gw) - 0.7
    
    rect = FancyBboxPatch((0.5, y-0.3), 11, 0.6, boxstyle="round,pad=0.05",
                           facecolor='#F3E5F5', edgecolor='#9C27B0', linewidth=1.2, alpha=0.5)
    ax.add_patch(rect)
    ax.text(6, y, 'Secure Data Communication Established', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#6A1B9A')
    
    y -= 0.7
    msg(y, 2.0, 10.0, 'PID | SE(K_GW-D, sensor_data)',
        'Encrypted data (48 B)', '#6A1B9A')
    
    y -= 0.55
    action(y, 10.0, 'Decrypt & process data', '#FFF3E0')
    
    # AS pseudonym rotation
    y -= 0.7
    action(y, 6.0, 'Rotate: PID_old←PID_curr, PID_curr←H(ID||m_new)', '#E3F2FD', width=4.0)
    y -= 0.55
    action(y, 6.0, 'm_old ← m_curr,  m_curr ← m_new', '#E3F2FD')
    
    ax.text(6, 1.3, '※ Session key K_GW-D enables direct device-to-gateway encrypted communication',
            ha='center', va='center', fontsize=8.5, style='italic', color='#666')
    
    plt.tight_layout()
    plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def draw_system_architecture(filepath):
    """Draw system architecture diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    
    ax.text(5, 6.7, 'System Architecture', fontsize=16, fontweight='bold',
            ha='center', va='center', color='#1a3c6e')
    
    # Gateway
    gw = FancyBboxPatch((3.8, 5.2), 2.4, 0.9, boxstyle="round,pad=0.15",
                         facecolor='#FF9800', edgecolor='#E65100', linewidth=2)
    ax.add_patch(gw)
    ax.text(5, 5.65, 'Gateway (GW)', ha='center', va='center', fontsize=11,
            fontweight='bold', color='white')
    ax.text(5, 5.35, 'RPL Root Node', ha='center', va='center', fontsize=8, color='white')
    
    # AS nodes
    for i, (x, label) in enumerate([(1.5, 'AS₁'), (5, 'AS₂'), (8.5, 'AS₃')]):
        as_box = FancyBboxPatch((x-1.0, 3.2), 2.0, 0.8, boxstyle="round,pad=0.1",
                                 facecolor='#2196F3', edgecolor='#1565C0', linewidth=1.5)
        ax.add_patch(as_box)
        ax.text(x, 3.6, f'Auth Server ({label})', ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')
        
        # Connection to GW
        ax.annotate('', xy=(5, 5.2), xytext=(x, 4.0),
                    arrowprops=dict(arrowstyle='<->', color='#666', lw=1.2,
                                   connectionstyle='arc3,rad=0'))
    
    # Device nodes
    device_positions = [(0.8, 1.3), (2.2, 1.3), (3.6, 1.3),
                        (5, 1.3), (6.4, 1.3), (7.8, 1.3), (9.2, 1.3)]
    device_labels = ['D₁', 'D₂', 'D₃', 'D₄', 'D₅', 'D₆', 'D₇']
    
    for pos, label in zip(device_positions, device_labels):
        x, y = pos
        d_box = FancyBboxPatch((x-0.55, y-0.3), 1.1, 0.6, boxstyle="round,pad=0.1",
                                facecolor='#4CAF50', edgecolor='#2E7D32', linewidth=1.2)
        ax.add_patch(d_box)
        ax.text(x, y, label, ha='center', va='center', fontsize=9,
                fontweight='bold', color='white')
    
    # Connections from devices to AS
    for i, (x, _) in enumerate(device_positions[:3]):
        ax.annotate('', xy=(1.5, 3.2), xytext=(x, 1.6),
                    arrowprops=dict(arrowstyle='<->', color='#aaa', lw=0.8,
                                   connectionstyle='arc3,rad=0'))
    for x, _ in device_positions[3:5]:
        ax.annotate('', xy=(5, 3.2), xytext=(x, 1.6),
                    arrowprops=dict(arrowstyle='<->', color='#aaa', lw=0.8))
    for x, _ in device_positions[5:]:
        ax.annotate('', xy=(8.5, 3.2), xytext=(x, 1.6),
                    arrowprops=dict(arrowstyle='<->', color='#aaa', lw=0.8))
    
    # Legend
    legend_y = 0.3
    ax.text(1.5, legend_y, '■ Gateway', fontsize=9, color='#FF9800', fontweight='bold')
    ax.text(3.8, legend_y, '■ Auth Servers (authorized by GW)', fontsize=9, color='#2196F3', fontweight='bold')
    ax.text(8.0, legend_y, '■ Devices', fontsize=9, color='#4CAF50', fontweight='bold')
    
    # K_GW_AS label
    ax.text(3.2, 4.5, 'K_GW-AS', fontsize=8, color='#E65100', fontweight='bold',
            rotation=30, ha='center')
    ax.text(5.0, 4.5, 'K_GW-AS', fontsize=8, color='#E65100', fontweight='bold',
            ha='center')
    ax.text(7.0, 4.5, 'K_GW-AS', fontsize=8, color='#E65100', fontweight='bold',
            rotation=-30, ha='center')
    
    # CoAP label
    ax.text(1.5, 2.5, 'CoAP', fontsize=7.5, color='#666', ha='center')
    ax.text(5, 2.5, 'CoAP', fontsize=7.5, color='#666', ha='center')
    ax.text(8.5, 2.5, 'CoAP', fontsize=7.5, color='#666', ha='center')
    
    plt.tight_layout()
    plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


# ═══════════════════════════════════════════════════════════════════
# MAIN DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════════

def generate_paper():
    # Generate diagrams
    arch_path = os.path.join(OUT_DIR, '_fig_architecture.png')
    enroll_path = os.path.join(OUT_DIR, '_fig_enrollment.png')
    auth_path = os.path.join(OUT_DIR, '_fig_authentication.png')
    kex_path = os.path.join(OUT_DIR, '_fig_key_exchange.png')
    
    print("Generating diagrams...")
    draw_system_architecture(arch_path)
    draw_enrollment_diagram(enroll_path)
    draw_authentication_diagram(auth_path)
    draw_key_exchange_diagram(kex_path)
    print("Diagrams generated.")
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ─── TITLE ───
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Lightweight and Decoupled Distributed PUF-Based\nAuthentication for Multihop IoT Networks')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.font.name = 'Calibri'
    title.paragraph_format.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Implementation Report — COOJA Simulation on Contiki-NG')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.paragraph_format.space_after = Pt(20)
    
    # ─── ABSTRACT ───
    add_heading_styled(doc, 'Abstract', level=1)
    add_para(doc, (
        'This report presents the implementation and simulation-based evaluation of a '
        'lightweight and decoupled distributed authentication scheme for resource-constrained '
        'multihop IoT networks. The proposed scheme utilizes Physically Unclonable Functions (PUF) '
        'for device-specific authentication, hash-based one-way accumulators for membership '
        'verification, and session-based random values for authentication mask updates. '
        'Authentication responsibilities are shared among both parent and non-parent IoT nodes '
        '(authentication servers), authorized by the gateway, preventing any single node from '
        'being overburdened. The scheme operates through three main phases: Node Enrollment, '
        'Node Authentication, and Session Key Exchange with Authentication Mask Update. '
        'The implementation is realized using CoAP over Contiki-NG OS and evaluated through '
        'COOJA simulation with a 100-node network deployment. Results demonstrate successful '
        'authentication of all 20 newly joined devices with energy consumption in the range of '
        '0.019–0.036 Joules per device, confirming the scheme\'s lightweight and efficient nature.'
    ), size=10.5, space_after=12)
    
    add_para(doc, (
        'Keywords: IoT Authentication, PUF, Hash-based Accumulator, Decoupled Distributed, '
        'Multi-hop Networks, CoAP, Contiki-NG, COOJA Simulation'
    ), italic=True, size=10, space_after=16)
    
    # ─── I. INTRODUCTION ───
    add_heading_styled(doc, 'I. Introduction', level=1)
    add_para(doc, (
        'The Internet of Things (IoT) connects millions of smart devices that collect and transmit '
        'data to centralized platforms such as gateways for decision making. However, IoT nodes are '
        'often deployed in adverse conditions, making authentication inevitable before communication. '
        'As these IoT nodes are battery-powered and resource-constrained, the authentication scheme '
        'should have a lightweight yet secure design.'
    ))
    add_para(doc, (
        'In the literature, distributed authentication schemes enable multiple entities in a multihop '
        'IoT network to perform authentication while distributing responsibilities across various '
        'points. However, schemes that assign parent IoT nodes as authentication servers may overwhelm '
        'them with recurring requests. This bottleneck can be mitigated by a decoupled distributed '
        'design, where non-parent IoT nodes share authentication duties with parent nodes — meaning '
        'any IoT node in the network can be authorized by the gateway to authenticate newly joined nodes.'
    ))
    add_para(doc, (
        'This report presents the implementation of such a decoupled distributed authentication '
        'scheme with the following key features:'
    ))
    
    features = [
        'PUF-based authentication eliminates the need for long-term storage of secret keys',
        'Hash-based one-way accumulator provides efficient membership verification',
        'Session-based random values prevent adversarial prediction through interception',
        'Lightweight operations (hash, XOR, two AES encryptions) minimize computational overhead',
        'Pseudonym-based communication preserves device anonymity',
        'Decoupled architecture distributes authentication load across the network',
    ]
    for f in features:
        add_bullet(doc, f)
    
    # ─── II. SYSTEM MODEL ───
    add_heading_styled(doc, 'II. System Model and Architecture', level=1)
    add_para(doc, (
        'The proposed scheme operates in a multihop IoT network consisting of three types of entities:'
    ))
    
    add_para(doc, 'Gateway (GW): The central trusted entity that serves as the RPL root node. '
             'It selects authentication servers, shares secret keys with them (K_GW-AS), and '
             'manages session keys with authenticated devices. Node ID: 1.', bold=False)
    add_para(doc, 'Authentication Server (AS): A resource-constrained IoT node authorized by the '
             'gateway to authenticate newly joined nodes. AS may be a parent or non-parent node, '
             'located at single-hop or multi-hop distance from the device being authenticated. '
             'It shares a pre-established secret key with the gateway (K_GW-AS) and the device '
             '(K_AS-D). Node ID: 2.', bold=False)
    add_para(doc, 'Device Node (D): A newly joined IoT node seeking to authenticate with the '
             'network. Each device is embedded with a PUF and possesses a unique private secret '
             'value y_D. Device nodes have IDs starting from 3 onward.', bold=False)
    
    # Architecture diagram
    doc.add_picture(arch_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Figure 1: System Architecture — Gateway, Authentication Servers, and Device Nodes',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── III. PREREQUISITES ───
    add_heading_styled(doc, 'III. Prerequisites', level=1)
    
    doc.add_heading('A. Physically Unclonable Function (PUF)', level=2)
    add_para(doc, (
        'A PUF operates as a one-way function that maps a challenge C to a unique response R, '
        'leveraging inherent hardware variations in the IoT device. The response is device-specific '
        'and infeasible to clone or replicate. Key properties of PUF include:'
    ))
    props = [
        'Evaluatable: A node with a legitimate PUF can efficiently compute R = PUF(C)',
        'Unpredictability: An attacker cannot guess R better than random without PUF access',
        'Uniqueness: Two different PUF instances never produce the same response for a given challenge',
        'Reliability: Same challenge always produces the same response on the same device',
        'One-Wayness: Inverting a response R to obtain the challenge C is infeasible',
    ]
    for p in props:
        add_bullet(doc, p, size=10)
    
    add_para(doc, (
        'In the implementation, PUF is simulated using a deterministic hash-based function that '
        'combines the node ID with the challenge through multiplicative hashing, ensuring unique '
        'responses per device per challenge while maintaining reproducibility.'
    ))
    
    doc.add_heading('B. Hash-based One-Way Accumulator', level=2)
    add_para(doc, (
        'A hash-based accumulator is a lightweight cryptographic construct where multiple elements '
        'are hashed and combined into a single compact value using bitwise AND operations. '
        'Individual elements can prove their membership within the accumulated value without '
        'revealing other values. The quasi-commutative property ensures the accumulation order '
        'does not affect the final value.'
    ))
    add_para(doc, (
        'Given a set Y = {y₁, y₂, ..., yₘ}, each element yᵢ is hashed: Yᵢ = H(yᵢ). '
        'These hashed values are accumulated: T_acc = T_acc & Y₁ & Y₂ & ... & Yₘ. '
        'Membership verification: if (T_acc & H(yᵢ)) == T_acc, then yᵢ is a member. '
        'The accumulator is initialized to 0xFF...FF (all ones), so each AND operation '
        'progressively sets bits to zero, and membership checks verify consistency.'
    ))
    
    # ─── IV. PROPOSED SCHEME ───
    add_heading_styled(doc, 'IV. Proposed Authentication Scheme', level=1)
    add_para(doc, (
        'The proposed scheme consists of three main phases: (1) Node Enrollment Phase, '
        '(2) Node Authentication Phase, and (3) Session Key Exchange and Authentication '
        'Mask Update Phase. Table I summarizes the notations used throughout the scheme.'
    ))
    
    # Notations Table
    notations = [
        ('Notation', 'Description'),
        ('H(.)', 'Collision-resistant SHA-256 hash function'),
        ('SE(.)/SD(.)', 'AES-128-ECB symmetric encryption/decryption'),
        ('PUF_D(.)', 'PUF function embedded in device D'),
        ('PUF_AS(.)', 'PUF function embedded in authentication server AS'),
        ('c_D', 'Challenge for device D issued by AS'),
        ('c_AS-D', 'Challenge for AS issued by device D'),
        ('R_D', 'PUF response: R_D = PUF_D(c_D)'),
        ('R_AS', 'PUF response: R_AS = PUF_AS(c_AS-D)'),
        ('y_D', 'Unique private secret of device D'),
        ('Y_D^H', 'Hash of secret: Y_D^H = H(y_D)'),
        ('m_D', 'Session-based random shared between D and AS'),
        ('T_acc', 'AND-based accumulated value maintained by AS'),
        ('Φ_AS-D', 'PUF binding: Φ_AS-D = R_AS ⊕ R_D'),
        ('K_GW-AS', 'Shared secret key between GW and AS'),
        ('K_AS-D', 'Shared secret key between AS and D (enrollment channel)'),
        ('K_GW-D', 'Session key between GW and D'),
        ('PID', 'Pseudonym: PID = H(ID_D || m_D)'),
        ('ts₁, ts₂', 'Timestamp components for freshness'),
        ('ts_auth', 'Authentication completion timestamp'),
    ]
    
    table = doc.add_table(rows=len(notations), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (n, d) in enumerate(notations):
        table.cell(i, 0).text = n
        table.cell(i, 1).text = d
        if i == 0:
            set_cell_shading(table.cell(i, 0), '1A3C6E')
            set_cell_shading(table.cell(i, 1), '1A3C6E')
            for p in table.cell(i, 0).paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                    r.font.size = Pt(10)
            for p in table.cell(i, 1).paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                    r.font.size = Pt(10)
        else:
            for p in table.cell(i, 0).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = 'Consolas'
            for p in table.cell(i, 1).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    
    add_para(doc, 'Table I: Notations Used in the Proposed Scheme',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── Phase A: Enrollment ───
    doc.add_heading('A. Node Enrollment Phase', level=2)
    add_para(doc, (
        'The Node Enrollment Phase is the first phase undergone by a newly joined device D with '
        'its designated authentication server AS. This phase is executed over a secure (encrypted) '
        'channel established between D and AS using their shared key K_AS-D. Every newly joined '
        'node executes this phase once until it gets disconnected from the network. The main '
        'purpose is to enroll the unique secret of D (y_D) with the accumulated value (T_acc) '
        'of AS, so that during authentication, D can prove its membership.'
    ))
    add_para(doc, 'The enrollment proceeds in two steps:', bold=True)
    
    add_para(doc, (
        'Step 1 — Registration Request (Reg-0): D initiates by sending its encrypted identity '
        'SE(K_AS-D, [ID_D | pad]) to AS (16 bytes). AS responds by generating a fresh random '
        'challenge c_D and session-based random m_D (32 bytes), sending them back encrypted: '
        'SE(K_AS-D, [c_D | m_D | pad]) as a 48-byte payload.'
    ))
    add_para(doc, (
        'Step 2 — Secret Enrollment (Reg-1): D computes PUF response R_D = PUF_D(c_D) using '
        'the received challenge. D generates its unique private secret y_D and computes '
        'Y_D^H = H(y_D). D also selects a challenge c_AS-D for AS. D sends the encrypted payload '
        'SE(K_AS-D, [ID_D | Y_D^H | R_D | c_AS-D | pad]) (48 bytes) to AS. Upon receiving, AS '
        'updates the accumulator: T_acc = T_acc & Y_D^H. AS computes R_AS = PUF_AS(c_AS-D) and '
        'stores Φ_AS-D = R_AS ⊕ R_D. Finally, AS computes the initial pseudonym '
        'PID_curr = H(ID_D || m_D) and stores all enrollment data per device.'
    ))
    
    # Enrollment diagram
    doc.add_picture(enroll_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Figure 2: Node Enrollment Phase — Sequence Diagram',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── Phase B: Authentication ───
    doc.add_heading('B. Node Authentication Phase', level=2)
    add_para(doc, (
        'The Node Authentication Phase proves the membership of D\'s unique secret y_D to AS '
        'through a secure yet lightweight protocol. Since authentication occurs over an open '
        'channel, y_D is communicated in a masked form using PUF responses and session-based '
        'randoms, preventing interception by adversaries.'
    ))
    add_para(doc, 'The authentication proceeds as follows:', bold=True)
    add_para(doc, (
        'Device Side: D regenerates its PUF response R_D = PUF_D(c_D) and computes its current '
        'pseudonym PID = H(ID_D || m_D). It computes Y_D^H = H(y_D) and constructs the '
        'authentication mask: mask = H(R_D || m_D || PID || ts₁). The masked hash Y_ASD is '
        'computed as Y_ASD = Y_D^H ⊕ mask. D sends {PID | Y_ASD | ts₁} (65 bytes, plaintext) to AS.'
    ))
    add_para(doc, (
        'AS Side: AS searches its client database for a matching PID (checking both PID_curr and '
        'PID_old for desynchronization recovery). After verifying ts₁ freshness (must be strictly '
        'ahead of last seen value), AS recovers R_D using the stored PUF binding: '
        'R_AS = PUF_AS(c_AS-D), R_D = Φ_AS-D ⊕ R_AS. AS recomputes the same mask and recovers '
        'Y_D^H = Y_ASD ⊕ mask. Finally, AS performs the membership test: T_new = T_acc & Y_D^H. '
        'If T_new equals T_acc, D is authenticated (the bitwise AND property of the accumulator '
        'guarantees that ANDing an already-accumulated value does not change the result).'
    ))
    
    # Authentication diagram
    doc.add_picture(auth_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Figure 3: Node Authentication Phase — Sequence Diagram',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── Phase C: Key Exchange ───
    doc.add_heading('C. Session Key Exchange and Mask Update Phase', level=2)
    add_para(doc, (
        'After successful authentication, the AS performs the key exchange phase to: '
        '(1) establish a session key K_GW-D between the gateway and the device, '
        '(2) securely inform the gateway about the authenticated device, and '
        '(3) update the session-based random (authentication mask) for the next authentication session.'
    ))
    add_para(doc, 'The key exchange proceeds as follows:', bold=True)
    
    add_para(doc, (
        'Step 1 — Key Generation: AS generates a random n₁ (32 bytes) and computes '
        'm_new = H(n₁). The session key is calculated as K_GW-D = H(R_D || m_new).'
    ))
    add_para(doc, (
        'Step 2 — Masked Random Delivery: To securely send m_new to D, AS computes '
        'mH_mask = H(Y_D^H || m_D || R_D || ID_AS || PID || ts₂) and produces '
        'm_H = m_new ⊕ mH_mask. AS sends {ACK | m_H | ts₂} (34 bytes) to D.'
    ))
    add_para(doc, (
        'Step 3 — Gateway Token: AS constructs an encrypted auth token: '
        'SE(K_GW-AS, [ID_D | ID_AS | ts_auth | pad]) || SE(K_GW-AS, K_GW-D[0:15]) || '
        'SE(K_GW-AS, K_GW-D[16:31]). This 48-byte encrypted token, along with the new PID '
        '(32 bytes) and ID_AS (1 byte), is forwarded to GW as an 81-byte CoAP POST to '
        '/test/auth_token.'
    ))
    add_para(doc, (
        'Step 4 — Device Recovery: D receives the reply, reconstructs mH_mask using its '
        'known values (Y_D^H, m_D, R_D, ID_AS, PID, ts₂), recovers m_new = m_H ⊕ mH_mask, '
        'and computes K_GW-D = H(R_D || m_new). D updates its session random: m_D ← m_new, '
        'and rotates its pseudonym: PID ← H(ID_D || m_new).'
    ))
    add_para(doc, (
        'Step 5 — Pseudonym Rotation: AS rotates pseudonyms: PID_old ← PID_curr, '
        'PID_curr ← H(ID_D || m_new). Session randoms are also rotated: m_old ← m_curr, '
        'm_curr ← m_new. This enables desynchronization recovery if a reply is lost.'
    ))
    
    # Key exchange diagram
    doc.add_picture(kex_path, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Figure 4: Session Key Exchange and Mask Update Phase — Sequence Diagram',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── V. IMPLEMENTATION DETAILS ───
    add_heading_styled(doc, 'V. Implementation Details', level=1)
    
    doc.add_heading('A. Software Platform', level=2)
    add_para(doc, (
        'The scheme is implemented on the Contiki-NG operating system, an open-source OS designed '
        'for resource-constrained IoT devices. The implementation uses the CoAP (Constrained '
        'Application Protocol) for all inter-node communication, with RPL (Routing Protocol for '
        'Low-Power and Lossy Networks) handling the routing.'
    ))
    
    impl_params = [
        ('Parameter', 'Value'),
        ('Operating System', 'Contiki-NG'),
        ('Mote Type', 'Cooja Mote'),
        ('Application Layer', 'CoAP (libcoap)'),
        ('Network Layer', 'RPL Lite'),
        ('MAC Layer', 'CSMA'),
        ('Encryption', 'AES-128-ECB'),
        ('Hashing', 'SHA-256'),
        ('PUF Simulation', 'Deterministic hash mixing (node_id ⊕ challenge)'),
        ('CoAP Max Payload', '128 bytes'),
        ('Freshness Window', '120 seconds'),
    ]
    
    table = doc.add_table(rows=len(impl_params), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (p, v) in enumerate(impl_params):
        table.cell(i, 0).text = p
        table.cell(i, 1).text = v
        if i == 0:
            set_cell_shading(table.cell(i, 0), '1A3C6E')
            set_cell_shading(table.cell(i, 1), '1A3C6E')
            for par in table.cell(i, 0).paragraphs:
                for r in par.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
            for par in table.cell(i, 1).paragraphs:
                for r in par.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
    
    add_para(doc, 'Table II: Implementation Parameters',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('B. Source Code Structure', level=2)
    files_desc = [
        ('File', 'Description', 'Lines'),
        ('device-node.c', 'IoT device: enrollment, authentication, data communication', '444'),
        ('as-node.c', 'Authentication server: handles Reg-0, Reg-1, Auth+KeyExchange', '619'),
        ('gw-node.c', 'Gateway: RPL root, processes auth tokens and encrypted data', '271'),
        ('aes.c / aes.h', 'AES-128 encryption/decryption implementation', '—'),
        ('sha256.c / sha256.h', 'SHA-256 hash function implementation', '—'),
        ('project-conf.h', 'Configuration defines (node IDs, MAC params, energest)', '35'),
        ('Makefile', 'Build system configuration for Contiki-NG', '—'),
    ]
    
    table = doc.add_table(rows=len(files_desc), cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (f, d, l) in enumerate(files_desc):
        table.cell(i, 0).text = f
        table.cell(i, 1).text = d
        table.cell(i, 2).text = l
        if i == 0:
            for j in range(3):
                set_cell_shading(table.cell(i, j), '1A3C6E')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
    
    add_para(doc, 'Table III: Source Code Structure',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('C. CoAP Endpoints and Message Sizes', level=2)
    add_para(doc, 'The following CoAP resources are defined for inter-node communication:')
    
    endpoints = [
        ('Endpoint', 'Method', 'Direction', 'Payload Size', 'Purpose'),
        ('/test/reg', 'GET', 'D → AS', '16 B (encrypted)', 'Registration request (Reg-0)'),
        ('/test/reg1', 'POST', 'D → AS', '48 B (encrypted)', 'Secret enrollment (Reg-1)'),
        ('/test/auth', 'POST', 'D → AS', '65 B (plaintext)', 'Authentication + Key Exchange'),
        ('/test/auth_token', 'POST', 'AS → GW', '81 B (partial enc)', 'Auth token delivery'),
        ('/test/data', 'POST', 'D → GW', '48 B (partial enc)', 'Encrypted sensor data'),
    ]
    
    table = doc.add_table(rows=len(endpoints), cols=5, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(endpoints):
        for j, val in enumerate(row_data):
            table.cell(i, j).text = val
            if i == 0:
                set_cell_shading(table.cell(i, j), '1A3C6E')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
                        r.font.size = Pt(9)
            else:
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
    
    add_para(doc, 'Table IV: CoAP Endpoints and Message Sizes',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── VI. SECURITY ANALYSIS ───
    add_heading_styled(doc, 'VI. Security Analysis', level=1)
    
    doc.add_heading('A. Resistance to Replay Attack', level=2)
    add_para(doc, (
        'The scheme employs multiple mechanisms against replay attacks. During enrollment, '
        'communications occur over a secure (encrypted) channel, preventing interception. '
        'In the authentication phase, timestamp-based freshness checks (ts₁ must be strictly '
        'ahead of the last accepted value) ensure that replayed messages with outdated timestamps '
        'are rejected. In the key exchange phase, ts₂ is a sequential counter generated by AS, '
        'and ts_auth is clock-based, both checked for freshness at the recipient. Session-based '
        'randoms (m_D) are rotated after each successful authentication, making older messages '
        'unverifiable.'
    ))
    
    doc.add_heading('B. Resistance to Man-in-the-Middle Attack', level=2)
    add_para(doc, (
        'The enrollment phase is resistant to MITM attacks as it uses a pre-established encrypted '
        'channel (K_AS-D). In the authentication and key exchange phases, the use of one-way '
        'collision-resistant hash functions (SHA-256), unique PUF responses (R_D), and session-based '
        'randoms (m_D) prevents adversaries from successfully modifying intercepted messages. '
        'Without access to R_D and the current m_D, any modification attempt will fail the '
        'membership verification at AS.'
    ))
    
    doc.add_heading('C. Device Anonymity via Pseudonyms', level=2)
    add_para(doc, (
        'The scheme provides device anonymity through rotating pseudonyms. Instead of using the '
        'real device ID in authentication messages, D computes PID = H(ID_D || m_D) and uses '
        'this pseudonym in all communications. After each successful authentication, the pseudonym '
        'is rotated: PID_new = H(ID_D || m_new). An adversary observing the network cannot link '
        'two sessions to the same device, as the pseudonym changes with each session random update. '
        'Even the gateway looks up sessions by PID, not by real device ID.'
    ))
    
    doc.add_heading('D. Desynchronization Recovery', level=2)
    add_para(doc, (
        'The scheme maintains both PID_curr and PID_old (along with m_curr and m_old) at the AS '
        'to handle desynchronization scenarios. If a device\'s last authentication reply was lost, '
        'the device will retry with its old pseudonym and session random. AS detects this by matching '
        'against PID_old and uses m_old to unmask the authentication message, enabling recovery '
        'without requiring re-enrollment.'
    ))
    
    # ─── VII. PERFORMANCE ANALYSIS ───
    add_heading_styled(doc, 'VII. Performance Analysis', level=1)
    
    doc.add_heading('A. Simulation Setup', level=2)
    add_para(doc, (
        'The performance evaluation is conducted using the COOJA simulator (Contiki-NG) in headless '
        'mode within a Docker container. The simulation deploys a 100-node network with the following '
        'configuration:'
    ))
    
    sim_params = [
        ('Parameter', 'Value'),
        ('Total Nodes', '100'),
        ('Gateway Nodes', '1 (Node ID 1)'),
        ('Authentication Server Nodes', '79 (Node IDs 2–80)'),
        ('Device Nodes (newly joined)', '20 (Node IDs 81–100)'),
        ('Network Topology', '10×10 grid, 30-unit spacing'),
        ('Propagation Model', 'UDGM (Distance Loss)'),
        ('Simulation Time', '1800 seconds'),
        ('Supply Voltage', '3.0 V'),
        ('CPU Current', '1.8 mA'),
        ('LPM Current', '0.0545 mA'),
        ('TX Current', '17.4 mA'),
        ('RX Current', '18.8 mA'),
    ]
    
    table = doc.add_table(rows=len(sim_params), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (p, v) in enumerate(sim_params):
        table.cell(i, 0).text = p
        table.cell(i, 1).text = v
        if i == 0:
            set_cell_shading(table.cell(i, 0), '1A3C6E')
            set_cell_shading(table.cell(i, 1), '1A3C6E')
            for par in table.cell(i, 0).paragraphs:
                for r in par.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
            for par in table.cell(i, 1).paragraphs:
                for r in par.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
    
    add_para(doc, 'Table V: Simulation Parameters',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('B. Protocol Summary', level=2)
    add_para(doc, (
        'The simulation results demonstrate complete protocol success: all 20 newly joined device '
        'nodes were successfully registered and authenticated through the AS (Node 2). '
        'A total of 347 data messages were confirmed by the gateway, with 18 initial rejections '
        'attributed to a race condition where device data arrives before the auth token is processed. '
        'These rejections are resolved automatically on the next retry.'
    ))
    
    summary = [
        ('Metric', 'Result'),
        ('Devices Registered', '20/20 (100%)'),
        ('Devices Authenticated', '20/20 (100%)'),
        ('Data Messages Confirmed', '347'),
        ('GW Initial Rejections', '18 (resolved on retry)'),
        ('Simulation Duration', '1800 s simulated / 30 s real-time'),
    ]
    
    table = doc.add_table(rows=len(summary), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (m, r) in enumerate(summary):
        table.cell(i, 0).text = m
        table.cell(i, 1).text = r
        if i == 0:
            set_cell_shading(table.cell(i, 0), '1A3C6E')
            set_cell_shading(table.cell(i, 1), '1A3C6E')
            for par in table.cell(i, 0).paragraphs:
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
            for par in table.cell(i, 1).paragraphs:
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
    
    add_para(doc, 'Table VI: Protocol Execution Summary',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('C. Energy Consumption Analysis', level=2)
    add_para(doc, (
        'Energy consumption is measured using the Contiki-NG Energest module, which tracks '
        'CPU active time, low-power mode (LPM) time, radio transmission (TX), and radio '
        'reception (RX) in clock ticks. These are converted to energy values using the Cooja '
        'mote current profiles. The measurement captures the complete authentication cycle: '
        'from the start of the auth+data block to its completion (including key exchange '
        'and one data transmission).'
    ))
    
    add_para(doc, (
        'Table VII shows the per-device authentication energy consumption for all 20 device nodes. '
        'Energy consumption ranges from 0.0187 J (Device 92) to 0.0365 J (Device 88), with '
        'variation attributed to network conditions (routing path length, congestion, retransmissions). '
        'The average energy consumption across all devices is approximately 0.025 J.'
    ))
    
    # Per-device results table
    device_data = [
        ('Device ID', 'CPU Time (s)', 'Energy (J)'),
        ('81', '0.4189', '0.0258'),
        ('82', '0.4321', '0.0266'),
        ('83', '0.3641', '0.0224'),
        ('84', '0.3451', '0.0213'),
        ('85', '0.3321', '0.0205'),
        ('86', '0.3561', '0.0219'),
        ('87', '0.3961', '0.0244'),
        ('88', '0.5913', '0.0365'),
        ('89', '0.5013', '0.0309'),
        ('90', '0.3223', '0.0199'),
        ('91', '0.3681', '0.0227'),
        ('92', '0.3040', '0.0187'),
        ('93', '0.3121', '0.0192'),
        ('94', '0.3330', '0.0205'),
        ('95', '0.4523', '0.0279'),
        ('96', '0.4453', '0.0275'),
        ('97', '0.4973', '0.0307'),
        ('98', '0.4963', '0.0306'),
        ('99', '0.5903', '0.0364'),
        ('100', '0.4373', '0.0270'),
    ]
    
    table = doc.add_table(rows=len(device_data), cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (did, cpu, energy) in enumerate(device_data):
        table.cell(i, 0).text = did
        table.cell(i, 1).text = cpu
        table.cell(i, 2).text = energy
        for j in range(3):
            for p in table.cell(i, j).paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)
            if i == 0:
                set_cell_shading(table.cell(i, j), '1A3C6E')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
    
    add_para(doc, 'Table VII: Per-Device Authentication Energy Consumption',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('D. AS Server Cumulative Energy', level=2)
    add_para(doc, (
        'The Authentication Server (Node 2) shows a cumulative CPU time of 208.51 seconds and '
        'total energy consumption of 12.88 Joules over the 1800-second simulation. This includes '
        'all enrollment (20 Reg-0 + 20 Reg-1), authentication (20 auth requests), key exchange '
        'operations, token deliveries to the gateway, and ongoing RPL routing duties. The energy '
        'cost per authentication (including key exchange and token forwarding) averages approximately '
        '0.64 J, which is dominated by radio reception (listening) during idle periods rather than '
        'the cryptographic operations themselves.'
    ))
    
    doc.add_heading('E. Computational Cost Comparison', level=2)
    add_para(doc, (
        'Table VIII compares the computational cost of the authentication and key exchange '
        'phase of the proposed scheme with related works. The proposed scheme achieves the lowest '
        'computational cost of 0.56 ms, utilizing 2 PUF operations, 8 SHA-256 hashes, '
        '2 AES encryptions, and 1 random generation — avoiding heavy ECC operations entirely.'
    ))
    
    comp_cost = [
        ('Scheme', 'Operations', 'Cost (ms)'),
        ('[8] Alruwaili et al.', '2T_puf + 8T_hash + 3T_ecc + 3T_rand', '3.39'),
        ('[9] Yang et al.', '6T_hash + 4T_ecc', '4.20'),
        ('[10] Li et al.', '2T_puf + 12T_hash + 2T_ecc + 4T_rand', '2.53'),
        ('[13] Zheng et al.', '2T_puf + 6T_hash + 4T_aes + 5T_rand', '0.67'),
        ('[6] Rosa et al. (LightSAE)', '14T_aes', '1.21'),
        ('Proposed Scheme', '2T_puf + 8T_hash + 2T_aes + 1T_rand', '0.56'),
    ]
    
    table = doc.add_table(rows=len(comp_cost), cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (s, o, c) in enumerate(comp_cost):
        table.cell(i, 0).text = s
        table.cell(i, 1).text = o
        table.cell(i, 2).text = c
        for j in range(3):
            for p in table.cell(i, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
            if i == 0:
                set_cell_shading(table.cell(i, j), '1A3C6E')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
        if i == len(comp_cost) - 1:
            for j in range(3):
                set_cell_shading(table.cell(i, j), 'E8F5E9')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.bold = True
    
    add_para(doc, 'Table VIII: Computational Cost Comparison (Authentication + Key Exchange)',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    doc.add_heading('F. Communication Cost Comparison', level=2)
    add_para(doc, (
        'Table IX compares the communication overhead. The proposed scheme uses three-message '
        'communication totaling only 928 bits — the lowest among all compared schemes. This is '
        'achieved by combining authentication and key exchange in a single protocol round and '
        'using compact message formats.'
    ))
    
    comm_cost = [
        ('Scheme', 'Messages', 'Cost (bits)'),
        ('[8] Alruwaili et al.', '2', '1408'),
        ('[9] Yang et al.', '2', '2612'),
        ('[10] Li et al.', '3', '4160'),
        ('[13] Zheng et al.', '3', '1600'),
        ('[6] Rosa et al. (LightSAE)', '6', '1536'),
        ('Proposed Scheme', '3', '928'),
    ]
    
    table = doc.add_table(rows=len(comm_cost), cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (s, m, c) in enumerate(comm_cost):
        table.cell(i, 0).text = s
        table.cell(i, 1).text = m
        table.cell(i, 2).text = c
        for j in range(3):
            for p in table.cell(i, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
            if i == 0:
                set_cell_shading(table.cell(i, j), '1A3C6E')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
        if i == len(comm_cost) - 1:
            for j in range(3):
                set_cell_shading(table.cell(i, j), 'E8F5E9')
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.bold = True
    
    add_para(doc, 'Table IX: Communication Cost Comparison (Authentication + Key Exchange)',
             italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # ─── VIII. CONCLUSION ───
    add_heading_styled(doc, 'VIII. Conclusion', level=1)
    add_para(doc, (
        'This report presented the implementation and simulation-based evaluation of a lightweight '
        'and decoupled distributed PUF-based authentication scheme for multihop IoT networks. '
        'The key contributions and findings are summarized as follows:'
    ))
    
    conclusions = [
        'The decoupled architecture successfully distributes authentication responsibilities '
        'across both parent and non-parent nodes, preventing gateway and parent node overloading.',
        
        'PUF-based authentication eliminates long-term key storage requirements while maintaining '
        'device-specific uniqueness through hardware-level challenge-response mechanisms.',
        
        'The hash-based one-way accumulator enables efficient O(1) membership verification '
        'using only bitwise AND operations, without needing to iterate over enrolled device lists.',
        
        'Session-based randoms with pseudonym rotation provide forward secrecy and device anonymity, '
        'with built-in desynchronization recovery through dual PID/random storage.',
        
        'COOJA simulation with 100 nodes demonstrates 100% authentication success rate for 20 '
        'newly joined devices, with per-device energy consumption of 0.019–0.036 J — confirming '
        'the scheme\'s suitability for resource-constrained IoT deployments.',
        
        'Computational cost of 0.56 ms (the lowest among compared schemes) is achieved by '
        'leveraging lightweight PUF and hash operations with only two AES encryptions, '
        'completely avoiding heavy ECC/pairing operations.',
        
        'Communication cost of 928 bits (3 messages) represents the most efficient message '
        'exchange among all compared schemes.',
    ]
    
    for c in conclusions:
        add_bullet(doc, c, size=10.5)
    
    add_para(doc, (
        'The implementation confirms that the proposed scheme achieves a practical balance between '
        'security and efficiency, making it well-suited for deployment in real-world '
        'resource-constrained multihop IoT networks.'
    ), space_after=16)
    
    # ─── REFERENCES ───
    add_heading_styled(doc, 'References', level=1)
    refs = [
        '[1] I. Zhou et al., "Internet of Things 2.0: Concepts, Applications, and Future Directions," IEEE Access, vol. 9, pp. 70961–71012, 2021.',
        '[2] B. Kang, D. Kim, and H. Choo, "Internet of Everything: A Large-Scale Autonomic IoT Gateway," IEEE Trans. Multi-Scale Computing Systems, vol. 3, no. 3, pp. 206–214, 2017.',
        '[3] N. Lajnef, A. Mami, and F. Derbel, "Applications of Wireless Sensor Networks and IoT Frameworks in Industry 4.0," Sensors, vol. 22, no. 6, 2022.',
        '[4] J. Granjal, E. Monteiro, and J. S. Silva, "Security in the Integration of Low-Power Wireless Sensor Networks with the Internet," Ad Hoc Networks, vol. 24, pp. 264–287, 2015.',
        '[5] W. Alnahari and M. T. Quasim, "Authentication of IoT Device and IoT Server Using Security Key," in 2021 ICOTEN, pp. 1–9, 2021.',
        '[6] P. Rosa, A. Souto, and J. Cecílio, "Light-SAE: A Lightweight Authentication Protocol for Large-Scale IoT Environments," IEEE Trans. Network and Service Management, vol. 20, no. 3, pp. 2428–2441, 2023.',
        '[7] M. Tiloca et al., "Security Considerations in the IP-based Internet of Things," Internet-Draft, IRTF, Aug. 2023.',
        '[8] O. Alruwaili et al., "RAAF-MEC: Reliable and Anonymous Authentication Framework for IoT-Enabled MEC Environment," Internet of Things, vol. 29, p. 101459, 2025.',
        '[9] W. Yang et al., "SAKMS: A Secure Authentication and Key Management Scheme for IETF 6TiSCH Industrial Wireless Networks," IEEE Trans. Network Science and Engineering, vol. 11, no. 3, pp. 3174–3188, 2024.',
        '[10] S. Li et al., "A Provably Secure and Practical PUF-Based End-to-End Mutual Authentication Protocol for IoT," IEEE Sensors J., vol. 21, no. 4, pp. 5487–5501, 2020.',
        '[11] Y. Zheng, Y. Cao, and C.-H. Chang, "A PUF-Based Data-Device Hash for Tampered Image Detection," IEEE Trans. Information Forensics and Security, vol. 15, pp. 620–634, 2019.',
        '[12] R. Amin et al., "Group Authentication Protocols for IoT: Challenges and Solutions," arXiv:1909.06371, Sept. 2019.',
        '[13] Y. Zheng et al., "PUF-Based Mutual Authentication and Key Exchange Protocol for Peer-to-Peer IoT Applications," IEEE TDSC, vol. 20, no. 4, pp. 3299–3316, 2022.',
        '[14] G. Selander, J. P. Mattsson, and F. Palombini, "Ephemeral Diffie-Hellman Over COSE (EDHOC)," RFC 9528, IETF, Mar. 2024.',
        '[15] S. Al-Riyami and K. G. Paterson, "Certificate-Based Public-Key Cryptography," Certificateless Cryptography Overview, 2003.',
        '[16] P. Zhang et al., "A Blockchain-Based Authentication Scheme for IoT-Enabled Maritime Systems," IEEE Trans. ITS, vol. 24, no. 2, pp. 2322–2331, 2023.',
        '[17] M. T. Al Ahmed et al., "Authentication-Chains: Blockchain-Inspired Lightweight Authentication for IoT Networks," Electronics, vol. 12, p. 867, 2023.',
        '[18] H. Sikarwar, D. Das, and S. Kalra, "Efficient Authentication Scheme Using Blockchain in IoT Devices," in AINA, vol. 1151, pp. 630–641, 2020.',
        '[19] H. Weerasena and P. Mishra, "Lightweight Multicast Authentication in NoC-Based SoCs," in 2024 ISQED, pp. 1–8, 2024.',
        '[20] B. Blanchet et al., "ProVerif 2.00: Automatic Cryptographic Protocol Verifier, User Manual and Tutorial," 2018.',
        '[21] D. Dolev and A. Yao, "On the Security of Public Key Protocols," IEEE Trans. Information Theory, vol. 29, no. 2, pp. 198–208, 2003.',
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(2)
    
    # Save
    output_path = os.path.join(OUT_DIR, 'Proposed_Scheme_Paper.docx')
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    
    # Cleanup temp diagram files
    for f in [arch_path, enroll_path, auth_path, kex_path]:
        try:
            os.remove(f)
        except:
            pass
    
    return output_path


if __name__ == '__main__':
    generate_paper()
